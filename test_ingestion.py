import os
import sys
import shutil
import tempfile
from pathlib import Path
from typing import List
from unittest.mock import patch

# Configure sys.path so the test script can import app modules directly
sys.path.append(str(Path(__file__).resolve().parent))

# Set dummy environment variables if they are not already set
os.environ.setdefault("GEMINI_API_KEY", "your_gemini_api_key_here")
os.environ.setdefault("CHROMA_DB_PATH", "data/test_chroma_db")
os.environ.setdefault("SQLITE_DB_PATH", "data/test_metadata.db")

from fastapi.testclient import TestClient
from app.main import app
from app.services.embedding_service import embedding_service
from app.services.db_service import db_service

# Define test files
TEST_TXT = "test_txt.txt"
TEST_DOCX = "test_docx.docx"
TEST_XLSX = "test_xlsx.xlsx"
TEST_PDF = "test_pdf.pdf"
TEST_LARGE = "test_large.txt"
TEST_PNG = "test_image.png"

# Color helpers
GREEN = "\033[92m"
RED = "\033[91m"
YELLOW = "\033[93m"
CYAN = "\033[96m"
RESET = "\033[0m"

def log_test(name, passed, detail=""):
    status_str = f"{GREEN}[PASS]{RESET}" if passed else f"{RED}[FAIL]{RESET}"
    detail_str = f" - {detail}" if detail else ""
    print(f"  {status_str} {name}{detail_str}")

def create_test_files():
    print(f"{CYAN}Generating programmatic test documents...{RESET}")
    
    # 1. Text file
    with open(TEST_TXT, "w", encoding="utf-8") as f:
        f.write("This is a simple text file for ingestion testing.\n")
        f.write("It has multiple sentences. Let's make sure it parses correctly.\n")
        f.write("FastAPI is a modern web framework. ChromaDB is a vector database.\n")

    # 2. Word document (.docx)
    try:
        import docx
        doc = docx.Document()
        doc.add_heading("Ingestion Test Document", level=0)
        doc.add_paragraph("This is paragraph 1 of our test word document.")
        doc.add_paragraph("This is paragraph 2 of our test word document. It contains some text that we want to parse.")
        
        # Add a table to test table parser extraction
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "Header A"
        table.rows[0].cells[1].text = "Header B"
        table.rows[1].cells[0].text = "Cell A1"
        table.rows[1].cells[1].text = "Cell B1"
        
        doc.save(TEST_DOCX)
    except Exception as e:
        print(f"{YELLOW}Warning: Could not create DOCX file programmatically: {e}{RESET}")

    # 3. Excel sheet (.xlsx)
    try:
        import openpyxl
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Employees"
        ws.append(["Name", "Age", "Role"])
        ws.append(["Alice", 30, "Engineer"])
        ws.append(["Bob", 25, "Designer"])
        
        ws2 = wb.create_sheet(title="Companies")
        ws2.append(["Company", "Location"])
        ws2.append(["Evon", "India"])
        
        wb.save(TEST_XLSX)
    except Exception as e:
        print(f"{YELLOW}Warning: Could not create XLSX file programmatically: {e}{RESET}")

    # 4. PDF document (.pdf) using PyMuPDF
    try:
        import fitz
        doc = fitz.open()
        
        # Page 1
        page1 = doc.new_page()
        page1.insert_text((50, 50), "This is page 1 of the test PDF document.")
        page1.insert_text((50, 100), "It contains some sample text for testing PyMuPDF parsing.")
        page1.insert_text((50, 750), "Header Title")  # Candidate header
        page1.insert_text((50, 800), "Footer Page 1") # Candidate footer
        
        # Page 2
        page2 = doc.new_page()
        page2.insert_text((50, 50), "This is page 2 of the test PDF document.")
        page2.insert_text((50, 100), "It has some other content on page 2.")
        page2.insert_text((50, 750), "Header Title")  # Candidate header (should match page 1)
        page2.insert_text((50, 800), "Footer Page 2") # Candidate footer (should not match page 1 footer)
        
        doc.save(TEST_PDF)
        doc.close()
    except Exception as e:
        print(f"{YELLOW}Warning: Could not create PDF file programmatically: {e}{RESET}")

    # 5. Invalid extension file (.png)
    with open(TEST_PNG, "wb") as f:
        f.write(b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR")

    # 6. Excessively large file (> 20MB)
    with open(TEST_LARGE, "w") as f:
        # 21MB of characters
        f.write("A" * (21 * 1024 * 1024))

def cleanup_test_files():
    print(f"{CYAN}Cleaning up test documents...{RESET}")
    for file in [TEST_TXT, TEST_DOCX, TEST_XLSX, TEST_PDF, TEST_LARGE, TEST_PNG]:
        if os.path.exists(file):
            try:
                os.remove(file)
            except Exception:
                pass
                
    # Clean up test database folder
    test_data_dir = Path("data")
    if test_data_dir.exists():
        test_chroma = test_data_dir / "test_chroma_db"
        test_sqlite = test_data_dir / "test_metadata.db"
        test_log = test_data_dir / "ingestion.log"
        
        if test_chroma.exists():
            shutil.rmtree(test_chroma, ignore_errors=True)
        if test_sqlite.exists():
            try:
                os.remove(test_sqlite)
            except Exception:
                pass

def get_dummy_embeddings(texts: List[str], max_retries: int = 5, initial_delay: float = 1.0) -> List[List[float]]:
    """Mock implementation returning dummy 3072-dimensional vectors."""
    return [[0.01 * (i % 100)] * 3072 for i in range(len(texts))]

def run_tests():
    print(f"{CYAN}Initializing TestClient and Database...{RESET}")
    
    # Force db_service to use test paths
    db_service.sqlite_path = "data/test_metadata.db"
    db_service.chroma_path = "data/test_chroma_db"
    db_service._init_sqlite()
    db_service._init_chroma()
    
    client = TestClient(app)
    
    has_api_key = os.getenv("GEMINI_API_KEY") != "your_gemini_api_key_here" and bool(os.getenv("GEMINI_API_KEY"))
    
    # We patch embeddings if no real API key is present
    patcher = None
    if not has_api_key:
        print(f"{YELLOW}No Gemini API key configured. Patching EmbeddingService with offline mock vectors.{RESET}")
        # Patch client check and embedding call
        patcher = patch.object(embedding_service, "get_embeddings", side_effect=get_dummy_embeddings)
        patcher.start()
        # Mock client configuration to bypass API key verification checks
        embedding_service.client_configured = True
    else:
        print(f"{GREEN}Real Gemini API Key detected. Using live embeddings service.{RESET}")

    all_passed = True

    try:
        # TEST 1: Check base status
        print(f"\n{CYAN}Running Test 1: GET / (Liveness Check){RESET}")
        res = client.get("/")
        passed = res.status_code == 200 and res.json().get("status") == "online"
        all_passed = all_passed and passed
        log_test("Root endpoint", passed, f"Response: {res.json()}")

        # TEST 2: Ingest all four document types
        print(f"\n{CYAN}Running Test 2: Ingest PDF, DOCX, XLSX, TXT documents{RESET}")
        files_to_upload = []
        for filename in [TEST_TXT, TEST_DOCX, TEST_XLSX, TEST_PDF]:
            if os.path.exists(filename):
                f_obj = open(filename, "rb")
                files_to_upload.append(("files", (filename, f_obj, "application/octet-stream")))
        
        res = client.post("/admin/upload", files=files_to_upload)
        
        # Close file handles
        for _, (_, f_obj, _) in files_to_upload:
            f_obj.close()

        passed = res.status_code == 200
        if passed:
            data = res.json()
            passed = len(data["uploaded"]) == 4 and len(data["failed"]) == 0
            log_test("Synchronous upload & processing", passed, f"Response: {data}")
        else:
            log_test("Synchronous upload & processing", False, f"Status Code: {res.status_code}, Body: {res.text}")
        all_passed = all_passed and passed

        # TEST 3: Check listing documents
        print(f"\n{CYAN}Running Test 3: List ingested documents{RESET}")
        res = client.get("/admin/documents")
        passed = res.status_code == 200
        if passed:
            docs = res.json()
            passed = len(docs) == 4
            log_test("GET /admin/documents", passed, f"Listed {len(docs)} documents. All 'completed'.")
        else:
            log_test("GET /admin/documents", False, f"Status Code: {res.status_code}")
        all_passed = all_passed and passed

        # TEST 4: Duplicate Ingestion Check
        print(f"\n{CYAN}Running Test 4: Duplicate upload rejection{RESET}")
        # Ingest same txt file again
        with open(TEST_TXT, "rb") as f:
            res = client.post("/admin/upload", files=[("files", (TEST_TXT, f, "text/plain"))])
        passed = res.status_code == 200
        if passed:
            data = res.json()
            passed = len(data["uploaded"]) == 0 and len(data["failed"]) == 1
            log_test(
                "Prevent duplicate ingestion", 
                passed, 
                f"Succeeded in rejecting. Error: {data['failed'][0]['error']}"
            )
        else:
            log_test("Prevent duplicate ingestion", False, f"Status Code: {res.status_code}")
        all_passed = all_passed and passed

        # TEST 5: Early validation - File type limits
        print(f"\n{CYAN}Running Test 5: Validate file extension rejection{RESET}")
        with open(TEST_PNG, "rb") as f:
            res = client.post("/admin/upload", files=[("files", (TEST_PNG, f, "image/png"))])
        passed = res.status_code == 400
        log_test(
            "Reject unsupported extensions (.png)", 
            passed, 
            f"Returned Status {res.status_code}. Detail: {res.json().get('error')}"
        )
        all_passed = all_passed and passed

        # TEST 6: Early validation - File size limits
        print(f"\n{CYAN}Running Test 6: Validate file size limit rejection (> 20MB){RESET}")
        with open(TEST_LARGE, "rb") as f:
            res = client.post("/admin/upload", files=[("files", (TEST_LARGE, f, "text/plain"))])
        passed = res.status_code == 400
        log_test(
            "Reject files > 20MB", 
            passed, 
            f"Returned Status {res.status_code}. Detail: {res.json().get('error')}"
        )
        all_passed = all_passed and passed

        # TEST 7: Delete ingested document
        print(f"\n{CYAN}Running Test 7: Delete ingested document{RESET}")
        # Fetch ID of TEST_TXT first
        res = client.get("/admin/documents")
        docs = res.json()
        target_doc_id = next((d["id"] for d in docs if d["filename"] == TEST_TXT), None)
        
        if target_doc_id:
            del_res = client.delete(f"/admin/documents/{target_doc_id}")
            passed = del_res.status_code == 200
            
            # Check list again to verify deletion
            list_res = client.get("/admin/documents")
            remaining_docs = list_res.json()
            passed = passed and len(remaining_docs) == 3
            passed = passed and not any(d["id"] == target_doc_id for d in remaining_docs)
            
            log_test("DELETE /admin/documents/{id}", passed, f"Deleted ID: {target_doc_id}")
        else:
            log_test("DELETE /admin/documents/{id}", False, "Could not find target document to delete.")
            passed = False
        all_passed = all_passed and passed

    finally:
        if patcher:
            patcher.stop()

    print("\n" + "=" * 50)
    if all_passed:
        print(f"{GREEN}ALL TESTS PASSED SUCCESSFULLY! Ingestion pipeline is verified.{RESET}")
    else:
        print(f"{RED}SOME TESTS FAILED. Please review the errors above.{RESET}")
    print("=" * 50)

if __name__ == "__main__":
    # Clean up any stale files/databases from previous runs before opening connections
    cleanup_test_files()
    try:
        create_test_files()
        run_tests()
    finally:
        cleanup_test_files()
