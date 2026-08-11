import re
import fitz  # PyMuPDF
import docx
import openpyxl
from typing import List, Dict, Any
from app.config import logger

class ParserService:
    def __init__(self):
        # Regex to match page numbers: e.g., "1", "Page 1", "1 of 10", "- 1 -", "Page 1 of 10"
        self.page_number_pattern = re.compile(
            r"^\s*(?:page\s+)?-?\s*\d+\s*-?(?:\s+of\s+\d+)?\s*$", 
            re.IGNORECASE
        )

    def parse_file(self, file_path: str, filename: str) -> List[Dict[str, Any]]:
        """
        Parses a file based on its extension.
        Returns a list of dicts: [{"page_number": int, "text": str}]
        """
        ext = filename.split(".")[-1].lower()
        logger.info("Parsing file %s (extension: %s)", filename, ext)

        if ext == "pdf":
            pages = self._parse_pdf(file_path)
        elif ext == "docx":
            pages = self._parse_docx(file_path)
        elif ext == "xlsx":
            pages = self._parse_xlsx(file_path)
        elif ext == "txt":
            pages = self._parse_txt(file_path)
        else:
            raise ValueError(f"Unsupported file extension: {ext}")

        # Clean pages
        cleaned_pages = self._clean_document_pages(pages)
        logger.info("Successfully extracted text from %s. Extracted %d pages/sheets.", filename, len(cleaned_pages))
        return cleaned_pages

    def _parse_pdf(self, file_path: str) -> List[Dict[str, Any]]:
        pages = []
        try:
            with fitz.open(file_path) as doc:
                for idx, page in enumerate(doc):
                    text = page.get_text()
                    pages.append({
                        "page_number": idx + 1,
                        "text": text
                    })
        except Exception as e:
            logger.error("Error parsing PDF %s: %s", file_path, e)
            raise e
        return pages

    def _parse_docx(self, file_path: str) -> List[Dict[str, Any]]:
        try:
            doc = docx.Document(file_path)
            text_parts = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                p_text = para.text.strip()
                if p_text:
                    text_parts.append(p_text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    # Filter empty cells and join
                    row_text = [t for t in row_text if t]
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            full_text = "\n\n".join(text_parts)
            return [{"page_number": 1, "text": full_text}]
        except Exception as e:
            logger.error("Error parsing DOCX %s: %s", file_path, e)
            raise e

    def _parse_xlsx(self, file_path: str) -> List[Dict[str, Any]]:
        pages = []
        try:
            wb = openpyxl.load_workbook(file_path, data_only=True)
            for idx, sheet_name in enumerate(wb.sheetnames):
                sheet = wb[sheet_name]
                sheet_lines = []
                
                for row in sheet.iter_rows(values_only=True):
                    # Filter out completely empty rows
                    if all(val is None for val in row):
                        continue
                    
                    row_str = [str(val).strip() if val is not None else "" for val in row]
                    # Check if there is actual content in the row
                    if any(row_str):
                        sheet_lines.append(" | ".join(row_str))
                
                sheet_text = f"Sheet: {sheet_name}\n" + "\n".join(sheet_lines)
                pages.append({
                    "page_number": idx + 1,
                    "text": sheet_text
                })
        except Exception as e:
            logger.error("Error parsing XLSX %s: %s", file_path, e)
            raise e
        return pages

    def _parse_txt(self, file_path: str) -> List[Dict[str, Any]]:
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
            return [{"page_number": 1, "text": content}]
        except Exception as e:
            logger.error("Error parsing TXT %s: %s", file_path, e)
            raise e

    def _clean_document_pages(self, pages: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """
        Cleans text across document pages:
        1. Identifies and strips common headers/footers (present on >= 50% of pages).
        2. Strips page numbers matching common formats on the first/last lines of each page.
        3. Cleans excessive whitespace.
        """
        if not pages:
            return pages

        # Step 1: Extract first and last lines for header/footer detection (only relevant for multi-page docs)
        header_candidates = []
        footer_candidates = []
        
        for p in pages:
            lines = [line.strip() for line in p["text"].split("\n") if line.strip()]
            header_candidates.append(lines[0] if lines else "")
            footer_candidates.append(lines[-1] if len(lines) > 1 else "")

        # Find headers and footers that repeat across >= 50% of pages (if document has >= 3 pages)
        common_header = None
        common_footer = None
        
        if len(pages) >= 3:
            header_counts = {}
            footer_counts = {}
            for h in header_candidates:
                if h:
                    header_counts[h] = header_counts.get(h, 0) + 1
            for f in footer_candidates:
                if f:
                    footer_counts[f] = footer_counts.get(f, 0) + 1
            
            threshold = len(pages) / 2
            
            for h, count in header_counts.items():
                if count >= threshold:
                    common_header = h
                    break
                    
            for f, count in footer_counts.items():
                if count >= threshold:
                    common_footer = f
                    break

        cleaned_pages = []
        for p in pages:
            lines = p["text"].split("\n")
            cleaned_lines = []
            
            for line in lines:
                cleaned_line = line.strip()
                if not cleaned_line:
                    continue
                
                # Check for recurring header/footer
                if common_header and cleaned_line == common_header:
                    continue
                if common_footer and cleaned_line == common_footer:
                    continue
                
                cleaned_lines.append(cleaned_line)

            # Step 2: Remove lines that look like page numbers from the first 2 and last 2 lines
            if len(cleaned_lines) > 0:
                # Check first line
                if self.page_number_pattern.match(cleaned_lines[0]):
                    cleaned_lines.pop(0)
            if len(cleaned_lines) > 0:
                # Check new first line (if there was one)
                if self.page_number_pattern.match(cleaned_lines[0]):
                    cleaned_lines.pop(0)
            if len(cleaned_lines) > 0:
                # Check last line
                if self.page_number_pattern.match(cleaned_lines[-1]):
                    cleaned_lines.pop()
            if len(cleaned_lines) > 0:
                # Check new last line (if there was one)
                if self.page_number_pattern.match(cleaned_lines[-1]):
                    cleaned_lines.pop()

            # Step 3: Re-assemble text and clean excessive whitespace
            joined_text = "\n".join(cleaned_lines)
            
            # Replace multiple consecutive spaces with a single space
            joined_text = re.sub(r"[ \t]+", " ", joined_text)
            
            # Replace 3 or more consecutive newlines with exactly 2 newlines (preserves paragraph breaks)
            joined_text = re.sub(r"\n{3,}", "\n\n", joined_text)
            
            cleaned_pages.append({
                "page_number": p["page_number"],
                "text": joined_text.strip()
            })

        return cleaned_pages

# Global parser service instance
parser_service = ParserService()
